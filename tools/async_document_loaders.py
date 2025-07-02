"""
Async document loaders for various file formats.
"""

import aiofiles
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import json
import yaml
import PyPDF2
from docx import Document
import pandas as pd
import logging
from abc import ABC, abstractmethod
from concurrent.futures import ThreadPoolExecutor

logger = logging.getLogger(__name__)


class AsyncBaseLoader(ABC):
    """Base class for async document loaders"""
    
    def __init__(self):
        self._executor = ThreadPoolExecutor(max_workers=4)
    
    @abstractmethod
    async def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Load and parse document asynchronously"""
        pass
    
    async def validate_file(self, file_path: Path) -> bool:
        """Validate file exists and is accessible"""
        return file_path.exists() and file_path.is_file()
    
    def __del__(self):
        """Cleanup executor"""
        if hasattr(self, '_executor'):
            self._executor.shutdown(wait=False)


class AsyncTextLoader(AsyncBaseLoader):
    """Async loader for text files"""
    
    async def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Load text file asynchronously"""
        try:
            file_path = Path(file_path)
            if not await self.validate_file(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            
            return {
                "type": "text",
                "content": content,
                "metadata": {
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "size": len(content)
                }
            }
        except Exception as e:
            logger.error(f"Error loading text file: {e}")
            raise


class AsyncJSONLoader(AsyncBaseLoader):
    """Async loader for JSON files"""
    
    async def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Load JSON file asynchronously"""
        try:
            file_path = Path(file_path)
            if not await self.validate_file(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                data = json.loads(content)
            
            return {
                "type": "json",
                "content": data,
                "metadata": {
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "keys": list(data.keys()) if isinstance(data, dict) else None,
                    "items": len(data) if isinstance(data, (list, dict)) else None
                }
            }
        except Exception as e:
            logger.error(f"Error loading JSON file: {e}")
            raise


class AsyncYAMLLoader(AsyncBaseLoader):
    """Async loader for YAML files"""
    
    async def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Load YAML file asynchronously"""
        try:
            file_path = Path(file_path)
            if not await self.validate_file(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                data = yaml.safe_load(content)
            
            return {
                "type": "yaml",
                "content": data,
                "metadata": {
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "keys": list(data.keys()) if isinstance(data, dict) else None
                }
            }
        except Exception as e:
            logger.error(f"Error loading YAML file: {e}")
            raise


class AsyncPDFLoader(AsyncBaseLoader):
    """Async loader for PDF files"""
    
    async def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Load PDF file asynchronously"""
        try:
            file_path = Path(file_path)
            if not await self.validate_file(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # PDF processing is CPU-intensive, run in executor
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                self._executor,
                self._load_pdf_sync,
                file_path
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error loading PDF file: {e}")
            raise
    
    def _load_pdf_sync(self, file_path: Path) -> Dict[str, Any]:
        """Synchronous PDF loading"""
        text_content = []
        metadata = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "pages": 0
        }
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata["pages"] = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append({
                        "page": page_num + 1,
                        "text": text
                    })
        
        return {
            "type": "pdf",
            "content": text_content,
            "metadata": metadata
        }


class AsyncDOCXLoader(AsyncBaseLoader):
    """Async loader for DOCX files"""
    
    async def load(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Load DOCX file asynchronously"""
        try:
            file_path = Path(file_path)
            if not await self.validate_file(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # DOCX processing is CPU-intensive, run in executor
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                self._executor,
                self._load_docx_sync,
                file_path
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error loading DOCX file: {e}")
            raise
    
    def _load_docx_sync(self, file_path: Path) -> Dict[str, Any]:
        """Synchronous DOCX loading"""
        doc = Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            "type": "docx",
            "content": {
                "text": "\n".join(content),
                "paragraphs": content,
                "tables": tables
            },
            "metadata": {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "paragraphs_count": len(content),
                "tables_count": len(tables)
            }
        }


class AsyncCSVLoader(AsyncBaseLoader):
    """Async loader for CSV files"""
    
    async def load(self, file_path: Union[str, Path], 
                  encoding: str = 'utf-8') -> Dict[str, Any]:
        """Load CSV file asynchronously"""
        try:
            file_path = Path(file_path)
            if not await self.validate_file(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Use pandas in executor for better performance
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                self._executor,
                self._load_csv_sync,
                file_path,
                encoding
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Error loading CSV file: {e}")
            raise
    
    def _load_csv_sync(self, file_path: Path, encoding: str) -> Dict[str, Any]:
        """Synchronous CSV loading"""
        df = pd.read_csv(file_path, encoding=encoding)
        
        return {
            "type": "csv",
            "content": df.to_dict('records'),
            "metadata": {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "rows": len(df),
                "columns": list(df.columns),
                "shape": df.shape
            }
        }


class AsyncDocumentLoader:
    """Main async document loader that delegates to specific loaders"""
    
    def __init__(self):
        self.loaders = {
            '.txt': AsyncTextLoader(),
            '.json': AsyncJSONLoader(),
            '.yaml': AsyncYAMLLoader(),
            '.yml': AsyncYAMLLoader(),
            '.pdf': AsyncPDFLoader(),
            '.docx': AsyncDOCXLoader(),
            '.csv': AsyncCSVLoader()
        }
    
    async def load_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Load document based on file extension"""
        try:
            file_path = Path(file_path)
            extension = file_path.suffix.lower()
            
            if extension not in self.loaders:
                raise ValueError(f"Unsupported file type: {extension}")
            
            loader = self.loaders[extension]
            return await loader.load(file_path)
            
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            raise
    
    async def load_batch(self, file_paths: List[Union[str, Path]], 
                        max_concurrent: int = 5) -> List[Dict[str, Any]]:
        """Load multiple documents concurrently"""
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def load_with_semaphore(file_path):
            async with semaphore:
                try:
                    return await self.load_document(file_path)
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {e}")
                    return {
                        "error": str(e),
                        "file_path": str(file_path)
                    }
        
        tasks = [load_with_semaphore(fp) for fp in file_paths]
        return await asyncio.gather(*tasks)
    
    async def load_directory(self, directory: Union[str, Path], 
                           pattern: str = "*",
                           recursive: bool = False) -> List[Dict[str, Any]]:
        """Load all matching documents from a directory"""
        directory = Path(directory)
        if not directory.is_dir():
            raise ValueError(f"Not a directory: {directory}")
        
        # Find matching files
        if recursive:
            file_paths = list(directory.rglob(pattern))
        else:
            file_paths = list(directory.glob(pattern))
        
        # Filter to supported extensions
        supported_files = []
        for fp in file_paths:
            if fp.is_file() and fp.suffix.lower() in self.loaders:
                supported_files.append(fp)
        
        logger.info(f"Found {len(supported_files)} supported files in {directory}")
        
        # Load all files
        return await self.load_batch(supported_files)


# Resume-specific async loader
class AsyncResumeLoader:
    """Specialized async loader for resume files"""
    
    def __init__(self):
        self.document_loader = AsyncDocumentLoader()
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    async def load_resume(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Load and parse resume file"""
        try:
            file_path = Path(file_path)
            
            if file_path.suffix.lower() not in self.supported_formats:
                raise ValueError(
                    f"Unsupported resume format: {file_path.suffix}. "
                    f"Supported formats: {', '.join(self.supported_formats)}"
                )
            
            # Load document
            doc_data = await self.document_loader.load_document(file_path)
            
            # Extract text content based on document type
            if doc_data["type"] == "pdf":
                text = "\n".join([page["text"] for page in doc_data["content"]])
            elif doc_data["type"] == "docx":
                text = doc_data["content"]["text"]
            else:
                text = doc_data["content"]
            
            # Basic resume parsing (extend as needed)
            resume_data = {
                "raw_text": text,
                "file_info": doc_data["metadata"],
                "parsed_sections": self._parse_resume_sections(text)
            }
            
            return resume_data
            
        except Exception as e:
            logger.error(f"Error loading resume: {e}")
            raise
    
    def _parse_resume_sections(self, text: str) -> Dict[str, Any]:
        """Basic resume section parsing"""
        # This is a simplified parser - extend based on requirements
        sections = {
            "contact": None,
            "summary": None,
            "experience": [],
            "education": [],
            "skills": []
        }
        
        # Simple keyword-based section detection
        lines = text.split('\n')
        current_section = None
        
        section_keywords = {
            "experience": ["experience", "employment", "work history"],
            "education": ["education", "academic"],
            "skills": ["skills", "technical skills", "competencies"],
            "summary": ["summary", "objective", "profile"]
        }
        
        # Parse sections (this is a basic implementation)
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line starts a new section
            for section, keywords in section_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    current_section = section
                    break
        
        return sections
    
    async def load_multiple_resumes(self, file_paths: List[Union[str, Path]]) -> List[Dict[str, Any]]:
        """Load multiple resume files concurrently"""
        tasks = [self.load_resume(fp) for fp in file_paths]
        return await asyncio.gather(*tasks, return_exceptions=True)